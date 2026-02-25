from docx import Document
import json

def extract_docx_data(file_path):
    doc = Document(file_path)
    data = {
        "text": [],
        "tables": []
    }
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            data["text"].append(para.text.strip())
            
    # Extract data from tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        data["tables"].append(table_data)
        
    return data

if __name__ == "__main__":
    report_path = "C:\\Users\\aicil\\.gemini\\antigravity\\scratch\\temp_report.docx"
    output_path = "C:\\Users\\aicil\\.gemini\\antigravity\\scratch\\report_data.json"
    
    try:
        extracted_data = extract_docx_data(report_path)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(extracted_data, f, ensure_ascii=False, indent=4)
        print(f"Data successfully extracted to {output_path}")
    except Exception as e:
        print(f"Error: {e}")
