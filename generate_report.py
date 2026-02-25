from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Paths
image_path = r"C:\Users\aicil\.gemini\antigravity\brain\338a7255-3ee4-4763-99cd-c0651dfc32bc\mapa_estatico_priorizacion.png"
output_path = r"C:\Users\aicil\.gemini\antigravity\brain\338a7255-3ee4-4763-99cd-c0651dfc32bc\Informe_Priorizacion_Sarampion_V2_Completo.docx"

# Data
data = [
    ["LAS JOYAS", "20", "Muy Alta (Rojo)"],
    ["VICTORIA DE DURANGO", "18", "Muy Alta (Rojo)"],
    ["LA GUAJOLOTA", "8", "Alta (Naranja)"],
    ["COLONIA HIDALGO", "3", "Media (Amarillo)"],
    ["LA JOYA", "2", "Media (Amarillo)"],
    ["GÓMEZ PALACIO", "1", "Baja (Verde)"],
    ["VICENTE GUERRERO", "1", "Baja (Verde)"],
    ["CIHUACORA", "1", "Baja (Verde)"]
]

doc = Document()

# Title
title = doc.add_heading('Informe de Priorización de Vacunación - Sarampión 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Fecha: 15 de febrero de 2026')
doc.add_paragraph('Este documento presenta la priorización de las brigadas de vacunación basada en el análisis de los 52 casos confirmados del brote de sarampión.')

# Table
doc.add_heading('Distribución de Casos y Prioridad por Localidad', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Localidad'
hdr_cells[1].text = 'Casos'
hdr_cells[2].text = 'Nivel de Prioridad'

for loc, cases, priority in data:
    row_cells = table.add_row().cells
    row_cells[0].text = loc
    row_cells[1].text = cases
    row_cells[2].text = priority

# Map
doc.add_heading('Visualización Geográfica del Brote', level=1)
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6.0))
    doc.add_paragraph('Mapa 1: Distribución de casos y niveles de prioridad en el territorio.').alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('[Error: No se pudo encontrar la imagen del mapa]')

# Age Data
age_data = [
    ["20-39 años", "13", "Crítica"],
    ["0-4 años", "11", "Crítica"],
    ["10-14 años", "9", "Alta"],
    ["5-9 años", "8", "Alta"],
    ["15-19 años", "7", "Alta"],
    ["40+ años", "4", "Media"]
]

doc.add_heading('Distribución de Casos y Prioridad por Edad', level=1)
table_age = doc.add_table(rows=1, cols=3)
table_age.style = 'Light Grid Accent 1'
hdr_age = table_age.rows[0].cells
hdr_age[0].text = 'Grupo de Edad'
hdr_age[1].text = 'Casos'
hdr_age[2].text = 'Prioridad'

for age_grp, cases, priority in age_data:
    row_cells = table_age.add_row().cells
    row_cells[0].text = age_grp
    row_cells[1].text = cases
    row_cells[2].text = priority

# Summary
doc.add_heading('Observaciones Críticas', level=1)
p = doc.add_paragraph()
p.add_run('Criterios de Priorización Geográfica: ').bold = True
p.add_run('Se define prioridad Muy Alta para localidades con 10+ casos (Las Joyas y Victoria de Durango). ')

p2 = doc.add_paragraph()
p2.add_run('Criterios de Priorización por Edad: ').bold = True
p2.add_run('Se observa una vulnerabilidad crítica en niños menores de 5 años y una alta incidencia en adultos de 20-39 años.')

doc.add_paragraph('Se recomienda la intervención inmediata en las zonas y grupos etarios marcados con prioridad crítica para contener la propagación.')

doc.save(output_path)
print(f"Document updated and saved to {output_path}")
