import docx
import os

def extract_to_txt(file_path, output_txt):
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return
    
    doc = docx.Document(file_path)
    
    with open(output_txt, 'w', encoding='utf-8') as f:
        f.write("--- FULL TEXT ---\n")
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
                
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"\nTable {i}:\n")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")

if __name__ == "__main__":
    path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CASOS NOTIFICADOS\CASOS CONFIRMADOS 2026\Informe_Estrategia_Vacunacion_Sarampion_Durango_2026.docx"
    extract_to_txt(path, "report_content.txt")
    print("Extraction complete. Content saved to report_content.txt")
