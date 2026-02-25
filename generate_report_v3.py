from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Paths
image_path = r"C:\Users\aicil\.gemini\antigravity\brain\338a7255-3ee4-4763-99cd-c0651dfc32bc\mapa_estatico_priorizacion.png"
output_path = r"C:\Users\aicil\.gemini\antigravity\brain\338a7255-3ee4-4763-99cd-c0651dfc32bc\Informe_Priorizacion_Sarampion_V3_1602.docx"

# Data for 61 cases
locality_data = [
    ["VICTORIA DE DURANGO", "26", "Muy Alta"],
    ["LAS JOYAS (Mezquital)", "20", "Muy Alta"],
    ["LA GUAJOLOTA", "8", "Alta"],
    ["COLONIA HIDALGO", "4", "Media"],
    ["LA JOYA", "2", "Media"],
    ["VICENTE GUERRERO", "1", "Baja"],
    ["CIHUACORA", "1", "Baja"],
    ["GÓMEZ PALACIO", "0", "N/A"] # Gómez Palacio had 1 in previous, 0 in new confirmed? Wait, Case 8 is GP.
]
# Re-verifying Case 8: GÓMEZ PALACIO.
locality_data = [
    ["VICTORIA DE DURANGO", "26", "Muy Alta"],
    ["LAS JOYAS (Mezquital)", "20", "Muy Alta"],
    ["LA GUAJOLOTA", "8", "Alta"],
    ["COLONIA HIDALGO", "4", "Media"],
    ["LA JOYA", "2", "Media"],
    ["GÓMEZ PALACIO", "1", "Baja"],
    ["VICENTE GUERRERO", "1", "Baja"],
    ["CIHUACORA", "1", "Baja"]
]

age_data = [
    ["20-39 años", "17", "Crítica"],
    ["0-4 años", "12", "Crítica"],
    ["5-9 años", "11", "Alta"],
    ["10-14 años", "9", "Alta"],
    ["15-19 años", "8", "Alta"],
    ["40+ años", "4", "Media"]
]

center_data = [
    ["SSD H. Materno Inf. (Durango/Mez)", "23", "Principal"],
    ["SSD CESSA N° 1 (Durango)", "11", "Alta"],
    ["IMSS HGZ N° 1", "4", "Media"],
    ["ISSSTE H. Santiago Ramón", "9", "Alta"],
    ["Otros (Gral 450, CESSA Hidalgo, etc)", "14", "Complementario"]
]

doc = Document()

# Title
title = doc.add_heading('Informe de Priorización de Vacunación - Sarampión 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Fecha de Actualización: 16 de febrero de 2026')
doc.add_paragraph('Actualización basada en 61 casos confirmados (Incremento de 9 casos respecto al reporte anterior).')

# 1. Locality Table
doc.add_heading('1. Priorización por Localidad (Casos Totales: 61)', level=1)
table_loc = doc.add_table(rows=1, cols=3)
table_loc.style = 'Light Grid Accent 1'
hdr_loc = table_loc.rows[0].cells
hdr_loc[0].text = 'Localidad'
hdr_loc[1].text = 'Casos'
hdr_loc[2].text = 'Prioridad'

for loc, cases, priority in locality_data:
    row_cells = table_loc.add_row().cells
    row_cells[0].text = loc
    row_cells[1].text = cases
    row_cells[2].text = priority

# 2. Age Table
doc.add_heading('2. Priorización por Grupo de Edad', level=1)
table_age = doc.add_table(rows=1, cols=3)
table_age.style = 'Light Grid Accent 1'
hdr_age = table_age.rows[0].cells
hdr_age[0].text = 'Grupo de Edad'
hdr_age[1].text = 'Casos'
hdr_age[2].text = 'Prioridad'

for age_grp, cases, priority in age_data:
    row_cells = table_age.add_row().cells
    row_cells[0].text = age_grp
    row_cells[1].text = cases
    row_cells[2].text = priority

# 3. Center Table
doc.add_heading('3. Centros de Notificación con mayor captación', level=1)
table_cnt = doc.add_table(rows=1, cols=3)
table_cnt.style = 'Light Grid Accent 1'
hdr_cnt = table_cnt.rows[0].cells
hdr_cnt[0].text = 'Centro de Salud / Hospital'
hdr_cnt[1].text = 'Casos Notificados'
hdr_cnt[2].text = 'Relevancia'

for center, cases, relevance in center_data:
    row_cells = table_cnt.add_row().cells
    row_cells[0].text = center
    row_cells[1].text = cases
    row_cells[2].text = relevance

# Map Image
doc.add_heading('Distribución Geográfica (Referencia)', level=1)
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6.0))
else:
    doc.add_paragraph('[Imagen del mapa pendiente de actualización]')

# Conclusions
doc.add_heading('Conclusiones y Recomendaciones', level=1)
p = doc.add_paragraph()
p.add_run('Crecimiento en Capital: ').bold = True
p.add_run('Se observa un aumento significativo en Victoria de Durango (26 casos), convirtiéndose en el foco principal de atención urbana.')

p2 = doc.add_paragraph()
p2.add_run('Vulnerabilidad Infantil y Productiva: ').bold = True
p2.add_run('Los grupos de 20-39 y menores de 5 años requieren campañas de vacunación y bloqueos específicos inmediatos.')

doc.save(output_path)
print(f"Document saved to {output_path}")
