import docx
import os

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return
    
    doc = docx.Document(file_path)
    
    print("--- FULL TEXT ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_data))

if __name__ == "__main__":
    path = r"c:\Users\aicil\OneDrive\Escritorio\PVU\SARAMPIÓN\CASOS NOTIFICADOS\CASOS CONFIRMADOS 2026\Informe_Estrategia_Vacunacion_Sarampion_Durango_2026.docx"
    read_docx(path)
